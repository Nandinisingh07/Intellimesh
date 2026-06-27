import os
import sys
from pathlib import Path
from loguru import logger

# Add project root to sys.path
project_root = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(project_root))

from backend.config import settings

def create_cybersecurity_policy_pdf():
    logger.info("Generating cybersecurity_policy.pdf using PyMuPDF (fitz)...")
    try:
        import fitz
        doc = fitz.open()
        
        # Page 1: Policy Introduction and Data Classification
        page1 = doc.new_page()
        text_p1 = (
            "INTELLIMESH CORPORATE CYBERSECURITY POLICY\n"
            "Document ID: SEC-POL-2026-V1\n"
            "Effective Date: June 1, 2026\n"
            "Classification: CONFIDENTIAL\n\n"
            "1. PURPOSE AND SCOPE\n"
            "This document establishes the security frameworks for all data processing systems within "
            "the IntelliMesh ecosystem. It applies to all employees, contractors, and systems accessing "
            "the corporate network environment.\n\n"
            "2. DATA CLASSIFICATION LEVELS\n"
            "Corporate information must be classified into one of the following four categories:\n"
            " - PUBLIC: Unrestricted information suitable for external dissemination.\n"
            " - INTERNAL: Non-public operational details restricted to active employees.\n"
            " - CONFIDENTIAL: Highly sensitive proprietary info, including source code, credentials, "
            "and personal identifiable information (PII). Requires encryption at rest.\n"
            " - RESTRICTED: Regulatory, financial, or executive planning data. Strict need-to-know access controls.\n\n"
            "All offline systems must enforce encryption on Confidential and Restricted storage mediums."
        )
        page1.insert_text((50, 50), text_p1, fontsize=11, fontname="helv")

        # Page 2: Access Control and Password Policies
        page2 = doc.new_page()
        text_p2 = (
            "INTELLIMESH CORPORATE CYBERSECURITY POLICY\n"
            "Document ID: SEC-POL-2026-V1\n"
            "Effective Date: June 1, 2026\n"
            "Classification: CONFIDENTIAL\n\n"
            "3. ACCESS CONTROL & IDENTITY MANAGEMENT\n"
            "Access to internal networks must follow the Principle of Least Privilege (PoLP). "
            "Role-Based Access Control (RBAC) must be implemented for all database administrators.\n\n"
            "4. PASSWORD & CREDENTIAL REQUIREMENTS\n"
            "All system passwords must meet the following cryptographic requirements:\n"
            " - Minimum length of 12 characters.\n"
            " - Must contain at least one uppercase letter, one lowercase letter, one number, and one special character.\n"
            " - Multi-Factor Authentication (MFA) must be active for all remote and local administrative sessions.\n"
            " - Passwords must be rotated every 90 days. Re-use of the past 5 passwords is forbidden.\n"
            " - Accounts will be locked for 30 minutes after 5 consecutive failed login attempts."
        )
        page2.insert_text((50, 50), text_p2, fontsize=11, fontname="helv")

        # Page 3: Incident Response and Auditing
        page3 = doc.new_page()
        text_p3 = (
            "INTELLIMESH CORPORATE CYBERSECURITY POLICY\n"
            "Document ID: SEC-POL-2026-V1\n"
            "Effective Date: June 1, 2026\n"
            "Classification: CONFIDENTIAL\n\n"
            "5. INCIDENT RESPONSE PROTOCOL\n"
            "In the event of a suspected data breach or unauthorized system access, the security operations "
            "center (SOC) must be notified within 15 minutes. The response workflow is as follows:\n"
            " 1. IDENTIFICATION: Verify and document the indicators of compromise.\n"
            " 2. CONTAINMENT: Isolate affected network segments and disable compromised credentials.\n"
            " 3. ERADICATION: Clean filesystems, run malware scans, and close system vulnerabilities.\n"
            " 4. RECOVERY: Restore verified clean backups and re-enable services.\n"
            " 5. LESSONS LEARNED: Publish an incident report within 72 hours.\n\n"
            "6. AUDIT TRAILS & LOG RETENTION\n"
            "System logs must capture all read/write activities on restricted databases. Logs must be cryptographically "
            "hashed and retained in read-only write-once-read-many (WORM) storage for a minimum of 365 days."
        )
        page3.insert_text((50, 50), text_p3, fontsize=11, fontname="helv")
        
        pdf_path = settings.SAMPLE_DOCS_DIR / "cybersecurity_policy.pdf"
        doc.save(str(pdf_path))
        doc.close()
        logger.info(f"Successfully generated {pdf_path}")
    except Exception as e:
        logger.error(f"Error creating PDF: {e}")

def create_ml_research_notes_docx():
    logger.info("Generating ml_research_notes.docx using python-docx...")
    try:
        import docx
        doc = docx.Document()
        doc.add_heading("ML Research Notes: Transformer Architecture Internals", level=0)
        
        doc.add_heading("1. Attention Mechanism Mathematics", level=1)
        doc.add_paragraph(
            "The fundamental component of modern transformer models is the Scaled Dot-Product Attention mechanism. "
            "Given query vectors Q, key vectors K, and value vectors V, attention is calculated using the following math equation:"
        )
        doc.add_paragraph("Attention(Q, K, V) = softmax( (Q * K^T) / sqrt(d_k) ) * V")
        doc.add_paragraph(
            "Here, d_k represents the dimension of the key vectors. The scaling factor 1/sqrt(d_k) prevents the dot products "
            "from growing too large in magnitude, which would push the softmax function into regions with extremely small gradients."
        )
        
        doc.add_heading("2. Multi-Head Attention", level=1)
        doc.add_paragraph(
            "Instead of performing a single attention function with queries, keys, and values, transformers project Q, K, and V "
            "multiple times with different, learned linear projections. This is known as Multi-Head Attention, allowing the model "
            "to attend to information from different representation subspaces at different positions simultaneously."
        )
        
        doc.add_heading("3. BERT vs GPT Architectures", level=1)
        doc.add_paragraph(
            "While both models use the transformer block, their architectures are tailored for different paradigms:\n"
            " - BERT (Bidirectional Encoder Representations from Transformers): Encoder-only model. Uses bidirectional attention "
            "to pre-train on masked language modeling and next sentence prediction. Ideal for classification, NER, and extraction.\n"
            " - GPT (Generative Pre-trained Transformer): Decoder-only model. Uses causal masked attention so that tokens can "
            "only attend to past positions. Ideal for autoregressive text generation, conversation, and creative writing."
        )
        
        doc.add_heading("4. Fine-Tuning Strategies: LoRA", level=1)
        doc.add_paragraph(
            "Full fine-tuning of large models becomes prohibitively expensive. Parameter-Efficient Fine-Tuning (PEFT) "
            "remedies this. Low-Rank Adaptation (LoRA) freezes the pre-trained model weights and injects trainable rank "
            "decomposition matrices into each layer of the Transformer architecture. For a weight matrix W of shape d x k, "
            "LoRA factorizes the weight update delta-W as B * A, where B is d x r and A is r x k, with rank r << min(d, k)."
        )
        
        docx_path = settings.SAMPLE_DOCS_DIR / "ml_research_notes.docx"
        doc.save(str(docx_path))
        logger.info(f"Successfully generated {docx_path}")
    except Exception as e:
        logger.error(f"Error creating DOCX: {e}")

def create_meeting_transcript_txt():
    logger.info("Generating meeting_transcript.txt...")
    try:
        content = (
            "INTELLIMESH TEAM MEETING TRANSCRIPT\n"
            "Date: June 5, 2026\n"
            "Topic: Q3 Product Roadmap & Latency Optimizations\n"
            "Participants: Alex (Product Manager), Priya (Machine Learning Lead), James (Backend Infra), Nandini (Frontend Lead)\n\n"
            "[00:01:15] Alex: Okay team, let's kick off. Our main goal is to review the Q3 roadmap for Project IntelliMesh and discuss the performance issues we've been seeing during ingestion and local LLM query generation.\n"
            "[00:02:40] Nandini: From the UI side, the drag-and-drop works smoothly, but users are complaining that there is a long silence when a document is processing. We need real-time status updates on the dashboard. I've designed a stage progress bar showing Extracting -> Chunking -> Embedding -> Indexed.\n"
            "[00:04:10] James: I can build a Server-Sent Events endpoint for uploads. That way, as python-docx or PyMuPDF extracts text, we send an SSE event, then send another when the LangChain text splitter chunking completes, and finally when ChromaDB indexing finishes.\n"
            "[00:06:22] Priya: That's perfect, James. On the model side, we are seeing high latency with Mistral-7B on standard consumer RAM. I recommend we set the default local LLM to Phi-2. It has only 2.7 billion parameters, runs in under 2GB RAM when quantized to 4-bits, and has a context window of 2048 tokens. I've tested it, and the response times are under 3 seconds on an average CPU.\n"
            "[00:09:55] Alex: That's a huge improvement. Let's make Phi-2 Q4_K_M GGUF the standard model. What about embeddings?\n"
            "[00:10:30] Priya: We will stick to all-MiniLM-L6-v2. It's tiny, extremely fast, and generates high-quality 384-dimensional dense vectors. ChromaDB handles it easily.\n"
            "[00:13:12] James: Regarding vector storage, I'm setting up ChromaDB PersistentClient to save files in the data directory. I also implemented MMR retrieval. It prevents redundant chunks from filling up the LLM's limited context window. The similarity threshold is set to 0.3. If similarity falls below this, we return a fallback offline message.\n"
            "[00:17:45] Nandini: I'll integrate Recharts on the dashboard page to show documents by modality (PDF, DOCX, Image, Audio) and total chunks. It will give the user a clear picture of their offline database size.\n"
            "[00:22:15] Alex: Awesome. What's the timeline for deployment?\n"
            "[00:24:50] James: I will have the API routers ready by tomorrow. Nandini can finish the Vite and Tailwind dashboard by Thursday. We should be ready for testing by Friday.\n"
            "[00:28:10] Alex: Excellent work everyone. Let's execute. Meeting adjourned."
        )
        txt_path = settings.SAMPLE_DOCS_DIR / "meeting_transcript.txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(content)
        logger.info(f"Successfully generated {txt_path}")
    except Exception as e:
        logger.error(f"Error creating TXT: {e}")

def create_network_diagram_png():
    logger.info("Generating network_diagram_notes.png using PIL...")
    try:
        from PIL import Image, ImageDraw, ImageFont
        
        # Create a new white image
        width, height = 800, 600
        img = Image.new("RGB", (width, height), color="white")
        draw = ImageDraw.Draw(img)
        
        # Draw borders
        draw.rectangle([10, 10, width-10, height-10], outline="black", width=3)
        
        # Title
        draw.text((20, 20), "Project IntelliMesh - Offline System Architecture Notes", fill="black")
        
        # Box 1: Firewall
        draw.rectangle([50, 100, 200, 150], outline="red", width=2)
        draw.text((60, 115), "[Firewall Box]", fill="black")
        
        # Box 2: Load Balancer
        draw.rectangle([250, 100, 400, 150], outline="blue", width=2)
        draw.text((260, 115), "[Load Balancer Box]", fill="black")
        
        # Box 3: App Servers
        draw.rectangle([450, 100, 600, 150], outline="green", width=2)
        draw.text((460, 115), "[App Servers Box]", fill="black")
        
        # Box 4: Redis Cache Layer
        draw.rectangle([50, 250, 200, 300], outline="purple", width=2)
        draw.text((60, 265), "[Redis Cache Layer]", fill="black")

        # Box 5: Database
        draw.rectangle([250, 250, 400, 300], outline="orange", width=2)
        draw.text((260, 265), "[Database Box]", fill="black")

        # Box 6: ChromaDB (Vector DB)
        draw.rectangle([450, 250, 600, 300], outline="brown", width=2)
        draw.text((460, 265), "[ChromaDB Vector DB]", fill="black")

        # Box 7: Local GGUF LLM
        draw.rectangle([250, 400, 400, 450], outline="magenta", width=2)
        draw.text((260, 415), "[Local GGUF LLM]", fill="black")

        # Draw connecting labels
        draw.text((50, 500), "Flow: Firewall -> Load Balancer -> App Servers -> Database", fill="black")
        draw.text((50, 520), "RAG Pipeline: User Query -> Embedder -> ChromaDB -> LLM -> Response", fill="black")
        
        png_path = settings.SAMPLE_DOCS_DIR / "network_diagram_notes.png"
        img.save(str(png_path))
        logger.info(f"Successfully generated {png_path}")
    except Exception as e:
        logger.error(f"Error creating PNG: {e}")

def main():
    logger.info("Initializing sample docs generator...")
    settings.create_dirs()
    create_cybersecurity_policy_pdf()
    create_ml_research_notes_docx()
    create_meeting_transcript_txt()
    create_network_diagram_png()
    logger.info("Finished sample docs generator!")

if __name__ == "__main__":
    main()

