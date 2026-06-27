import docx
from typing import List, Dict, Any
from loguru import logger

class DocxProcessor:
    def __init__(self) -> None:
        pass

    def extract_text(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract text from a DOCX file using python-docx.
        Returns a list containing a single dict with the full text and page=1.
        """
        logger.info(f"Extracting text from DOCX: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            paragraphs = []
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text.strip())
                    
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            full_text = "\n\n".join(paragraphs)
            logger.info(f"Successfully extracted text from DOCX: {file_path}")
            return [{"text": full_text, "page": 1}]
        except Exception as e:
            logger.error(f"Error extracting DOCX text from {file_path}: {e}")
            raise e
